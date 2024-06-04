from flask import Flask, render_template, redirect, url_for, request, session, flash, Blueprint, send_from_directory, jsonify, abort
import json
import random
import string
import os
#from .models import Account
from huggingface_hub import InferenceClient
from docx import Document
import pypandoc
import fitz
import magic
import csv
import time
import pandas as pd
import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('Agg')
import io
import base64
from matplotlib.ticker import MaxNLocator

with open('/Users/susannamau/Dev/BPER/Challenge/config.json', 'r') as config_file:
    config = json.load(config_file)

if not os.path.exists(config["FEEDBACK_FILE"]):
    with open(config["FEEDBACK_FILE"], mode='w', newline='') as file:
        writer = csv.writer(file)
        writer.writerow(["File1","File2","File1 Word Count", "File2 Word Count", "Response Word Count", "Execution Time", "Rating"])


main = Blueprint('main', __name__)

@main.route('/')
def welcome():
    return render_template("task_choice.html")


@main.route('/task_choice', methods=['POST', 'GET'])
def task_choice():
    print(request.form['submit_button'])
    if request.form['submit_button'] == 'Question Answering':
        files = get_files()
        return render_template("quest_ans.html", files=files)
    elif request.form['submit_button'] == 'Documents Differences Finder':
        return render_template("doc_diffs.html")
    else:
        return "Invalid task selected", 400

@main.route('/quest_ans', methods=['POST', 'GET'])
def quest_ans():
    files = get_files()
    return render_template("quest_ans.html", files=files)


def get_files():
    files = os.listdir(config['UPLOAD_FOLDER'])
    files = [f for f in files if os.path.isfile(os.path.join(config['UPLOAD_FOLDER'], f))]
    return files



def allowed_file(filename):
    if filename.rsplit('.', 1)[1].lower() in config["ALLOWED_EXTENSIONS"]:
        return True
    else:
        return False
    


@main.route('/upload_qa', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return "Nessun file selezionato", 400

    file = request.files['file']
    if file.filename == '':
        return "Nome file non valido", 400

    if file and allowed_file(file.filename):
        os.makedirs(config['UPLOAD_FOLDER'], exist_ok=True)
        file_path = os.path.join(config['UPLOAD_FOLDER'], file.filename)
        file.save(file_path)
        files = get_files()
        return render_template("quest_ans.html", files=files)
    
    if allowed_file(file.filename) == False:
        return "Estensione file non valida", 400
    

    
@main.route('/delete_file', methods=['POST'])
def delete_file():
    data = request.get_json()
    filename = data['filename']
    file_path = os.path.join(config['UPLOAD_FOLDER'], filename)
    if os.path.exists(file_path):
        os.remove(file_path)
        return jsonify(success=True)
    return jsonify(success=False)



@main.route('/<filename>', methods=['POST', 'GET'])
def download_file(filename):
    file_path = os.path.join(config['UPLOAD_FOLDER'], filename)

    print(f"Verifica percorso file: {file_path}") 
    if not os.path.exists(file_path):
        print(f"File non trovato: {file_path}")

    return send_from_directory(config['UPLOAD_FOLDER'], filename)



def get_files_content(folder):
    files_content = []
    files = os.listdir(folder)
    for file in [f for f in files if os.path.isfile(os.path.join(folder, f))]:
        print("File: ", f"{folder}/{file}")
        ext = file.rsplit('.', 1)[1].lower()
        if ext == "txt":
            with open(f'{folder}/{file}', 'r') as file:
                content = file.read()
        elif ext == "csv":
            content = list(pd.read_csv(f'{folder}/{file}'))
        elif ext == "docx":
            doc = Document(f'{folder}/{file}')
            content = ""
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"
        elif ext == "doc":
            content = pypandoc.convert_file(f'{folder}/{file}', to="plain")
        elif ext == "pdf":
            pdf_document = fitz.open(f"{folder}/{file}")
            content = ""
            for page_num in range(len(pdf_document)):
                page = pdf_document.load_page(page_num)
                content += page.get_text()
        else:
            content = magic.Magic(mime=True)
            content = content.from_file(f"{folder}/{file}")
        files_content.append(content)
    return files_content


@main.route('/ask', methods=['POST'])
def ask_question():
    files_content = get_files_content(config['UPLOAD_FOLDER'])

    data = request.get_json()
    question = data.get('question')

    hf_token = config["HF_TOKEN"]
    hf_client = InferenceClient(token=hf_token)
    llm_model = config["LLM"]

    messages=[
    {"role": "system", "content": "Sei un assistente che risponde in Italiano alle domande dell'utente consultando i file a disposizione."},
    {"role": "system", "content": f"Date le seguenti informazioni: {files_content}"},
    {"role": "user", "content": question}
    ]

    completion = hf_client.chat_completion(model=llm_model, messages=messages, max_tokens=500)
    response = completion.choices[0].message.content
    
    if response != None:
        return jsonify({'answer': response})
    else:
        return jsonify({'answer': 'There was an error processing your question.'})
    
##############################################################
    
@main.route('/upload_diff', methods=['POST'])
def upload_file_diff():
    if 'file1' in request.files:
        file1 = request.files['file1']
        if file1.filename == '':
            return jsonify({'error': 'Nome file non valido'}), 400
        if file1 and allowed_file(file1.filename):
            os.makedirs(config['COMP_FOLDER'], exist_ok=True)
            file_path1 = os.path.join(config['COMP_FOLDER'], file1.filename)
            file1.save(file_path1)
            return jsonify({'filename': file1.filename})

    if 'file2' in request.files:
        file2 = request.files['file2']
        if file2.filename == '':
            return jsonify({'error': 'Nome file non valido'}), 400
        if file2 and allowed_file(file2.filename):
            os.makedirs(config['COMP_FOLDER'], exist_ok=True)
            file_path2 = os.path.join(config['COMP_FOLDER'], file2.filename)
            file2.save(file_path2)
            return jsonify({'filename': file2.filename})

    return jsonify({'error': 'Nessun file selezionato'}), 400

def get_file_content(file_path):
    file = file_path
    ext = file.rsplit('.', 1)[1].lower()
    if ext == "txt":
        with open(f'{file}', 'r') as file:
            content = file.read()
    elif ext == "csv":
        content = list(pd.read_csv(f'{file}'))
    elif ext == "docx":
        doc = Document(f'{file}')
        content = ""
        for paragraph in doc.paragraphs:
            content += paragraph.text + "\n"
    elif ext == "doc":
        content = pypandoc.convert_file(f'{file}', to="plain")
    elif ext == "pdf":
        pdf_document = fitz.open(f"{file}")
        content = ""
        for page_num in range(len(pdf_document)):
            page = pdf_document.load_page(page_num)
            content += page.get_text()
    else:
        content = magic.Magic(mime=True)
        content = content.from_file(f"{file}")
    return content

@main.route('/doc_diffs', methods=['POST'])
def doc_diffs():
    start_time = time.time()
    data = request.get_json()
    print(data)
    filename1 = data.get('filename1')
    filename2 = data.get('filename2')

    if not filename1 or not filename2:
        return jsonify({'error': 'Invalid input data.'}), 400
    
    file_path1 = os.path.join(config['COMP_FOLDER'], filename1)
    file_path2 = os.path.join(config['COMP_FOLDER'], filename2)
    #print(file_path1, file_path2)

    if not os.path.exists(file_path1) or not os.path.exists(file_path2):
        return jsonify({'error': 'One or both files do not exist.'}), 404

    try:
        ext1 = filename1.rsplit('.', 1)[1].lower()
        ext2 = filename2.rsplit('.', 1)[1].lower()

        content1 = get_file_content(file_path1)
        content2 = get_file_content(file_path2)

        len_input_1 = len(content1.split())
        len_input_2 = len(content2.split())

        hf_token = config["HF_TOKEN"]
        hf_client = InferenceClient(token=hf_token)
        llm_model = config["LLM"]

        messages = [
            {"role": "system", "content": "Sei un assistente intelligente incaricato di trovare le differenze tra documenti di testo e di rispondermi in italiano."},
            {"role": "system", "content": "Ti fornirò due documenti di testo e tu dovrai restituirmi una lista ordinata delle differenze tra di essi. Voglio che la risposta sia in italiano. Voglio che i titoli di ogni sezione della risposta siano tra ** e **."},
            {"role": "system", "content": f"Il primo documento è:\n\n{content1}"},
            {"role": "system", "content": f"Il secondo documento è:\n\n{content2}"},
            {"role": "user", "content": "Analizza i due documenti e forniscimi una lista dettagliata delle differenze tra di essi. Includi differenze di contenuto. Alla fine di tutte le differenze, scrivi un paragrafo intitolato Explainability dove spieghi come sei arrivato alla risposta che mi hai dato."}
        ]

        os.remove(file_path1)
        os.remove(file_path2)

        completion = hf_client.chat_completion(model=llm_model, messages=messages, max_tokens=1000)
        response = completion.choices[0].message.content
        #print(response)

        if response:
            session['file1'] = filename1
            session['file2'] = filename2
            session['file1_word_count'] = len_input_1
            session['file2_word_count'] = len_input_2
            session['response_word_count'] = len(response.split())
            session['execution_time'] = time.time() - start_time
            return jsonify({'differences': response})
        else:
            return jsonify({'error': 'There was an error processing your request.'})
        print("fine try")
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    

##############################################################

if not os.path.exists(config["FEEDBACK_FILE"]):
    with open(config["FEEDBACK_FILE"], mode='w', newline='') as file:
        writer = csv.writer(file)
        writer.writerow(["Rating"])

@main.route('/submit_feedback', methods=['POST', 'GET'])
def submit_feedback():
    data = request.get_json()
    rating = data.get('rating')

    if not rating:
        return jsonify({'success': False, 'error': 'Invalid rating'}), 400
    
    print(session)
    file1 = session.get('file1', '')
    file2 = session.get('file2', '')
    file1_word_count = session.get('file1_word_count', 0)
    file2_word_count = session.get('file2_word_count', 0)
    response_word_count = session.get('response_word_count', 0)
    execution_time = session.get('execution_time', 0)

    try:
        with open(config["FEEDBACK_FILE"], mode='a', newline='') as file:
            writer = csv.writer(file)
            writer.writerow([file1, file2, file1_word_count, file2_word_count, response_word_count, execution_time, rating])
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500
    

##################################

@main.route('/login', methods=['POST', 'GET'])
def admin_login():
    return render_template('login.html')

@main.route('/admin-dashboard', methods=['POST', 'GET'])
def admin_dashboard():
    username = request.form.get('username')
    password = request.form.get('password')
    if username == config['ADMIN_USERNAME'] and password == config['ADMIN_PASSWORD']:
        df = pd.read_csv(config['FEEDBACK_FILE'])
        avg_rating = round(df['Rating'].mean(), 2)
        avg_response_time = round(df['Execution Time'].mean(), 2)

        # File con rating bassi
        df["Execution Time"] = df["Execution Time"].round(2)
        low_ratings_df = df[df['Rating'].isin([1, 2])]
        
        # Genera il grafico del tempo in funzione del numero di parole
        plt.figure(figsize=(10, 6))
        plt.scatter(df['Response Word Count'], df['Execution Time'], alpha=0.5)
        plt.title('Response Time vs Number of Words of Response')
        plt.xlabel('Response Word Count')
        plt.ylabel('Response Time (s)')
        scatter_img = io.BytesIO()
        plt.savefig(scatter_img, format='png')
        scatter_img.seek(0)
        scatter_url = base64.b64encode(scatter_img.getvalue()).decode('utf8')
        plt.close()  # Chiude la figura per liberare memoria
        
        # Genera l'istogramma dei rating
        plt.figure(figsize=(10, 6))
        df['Rating'].hist(bins=[1, 2, 3, 4, 5, 6], edgecolor='black', align='left')
        plt.title('Distribution of Ratings')
        plt.xlabel('Rating')
        plt.ylabel('Frequency')
        plt.xticks([1, 2, 3, 4, 5])
        plt.gca().yaxis.set_major_locator(MaxNLocator(integer=True))
        hist_img = io.BytesIO()
        plt.savefig(hist_img, format='png')
        hist_img.seek(0)
        hist_url = base64.b64encode(hist_img.getvalue()).decode('utf8')
        plt.close()  # Chiude la figura per liberare memoria
        
        return render_template('admin_dashboard.html', avg_rating=avg_rating, avg_response_time=avg_response_time,
                               scatter_url=scatter_url, hist_url=hist_url, low_ratings_df=low_ratings_df.to_dict(orient='records'))
    else:
        return "Invalid credentials", 401
